"""สร้าง Project Documentation Word (.docx) สำหรับระบบบัญชี AccCloud."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import date

TODAY = date.today().strftime("%d/%m/%Y")

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── Styles helpers ──────────────────────────────────────────────────────────────
ACCENT  = RGBColor(0x1D, 0x4E, 0xD8)   # blue-700
ACCENT2 = RGBColor(0x06, 0x95, 0x5F)   # green-700
GRAY    = RGBColor(0x6B, 0x72, 0x80)
BLACK   = RGBColor(0x11, 0x18, 0x27)
LIGHT   = RGBColor(0xE5, 0xE7, 0xEB)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = ACCENT if level <= 2 else BLACK
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.bold = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(text: str = "", bold=False, size=10, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.75 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_table(headers: list, rows: list, col_widths: list = None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, "1D4ED8")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # Data rows
    for ri, row in enumerate(rows):
        trow = t.rows[ri + 1]
        if ri % 2 == 1:
            for cell in trow.cells:
                set_cell_bg(cell, "F3F4F6")
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("AccCloud")
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = ACCENT

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("ระบบบัญชีออนไลน์สำหรับสำนักงานบัญชี")
r2.font.size  = Pt(16)
r2.font.color.rgb = GRAY

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Project Documentation")
r3.font.size  = Pt(22)
r3.font.bold  = True

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(f"Version 1.0  |  {TODAY}")
r4.font.size  = Pt(11)
r4.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════════════════════════════════════════════
add_heading("1. ภาพรวมระบบ (System Overview)", 1)
add_para(
    "AccCloud เป็นระบบบัญชีออนไลน์แบบ Multi-tenant สำหรับสำนักงานบัญชี (Accounting Firm) "
    "ที่ต้องบริหารจัดการบัญชีลูกค้าหลายบริษัทพร้อมกัน รองรับมาตรฐาน NPAEs ของไทย "
    "ทำงานบน FastAPI + SQLite (dev) / PostgreSQL (prod) และแสดงผล UI ด้วย HTMX + Alpine.js + Tailwind CSS",
    size=10
)

doc.add_paragraph()
add_heading("Tech Stack", 2)
add_table(
    ["Layer", "Technology"],
    [
        ["Backend Framework", "Python 3.11+ / FastAPI (async)"],
        ["ORM", "SQLAlchemy 2.0 (Mapped[], mapped_column())"],
        ["Database (Dev)", "SQLite — แยก DB ต่อบริษัท"],
        ["Database (Prod)", "PostgreSQL"],
        ["Frontend", "Jinja2 Templates + HTMX + Alpine.js"],
        ["CSS", "Tailwind CSS (CDN) + theme.css design system"],
        ["Auth", "JWT (httpOnly cookie) + Role-based"],
        ["OCR", "Anthropic Claude Vision API (claude-sonnet-4-6)"],
        ["PDF Export", "WeasyPrint"],
        ["Excel Export", "openpyxl"],
        ["Deploy", "Docker + Docker Compose + Nginx"],
        ["Migration", "Alembic (shared DB) + scripts/migrate_company_db.py (company DB)"],
    ],
    [5, 11]
)

# ════════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════════
add_heading("2. สถาปัตยกรรมระบบ (Architecture)", 1)

add_heading("2.1 Multi-tenant Structure", 2)
add_para("ระบบแบ่งข้อมูลออกเป็น 3 ระดับ:", size=10)
add_bullet("Firm (สำนักงานบัญชี) — ระดับสูงสุด")
add_bullet("Company (บริษัทลูกค้า) — แต่ละ Firm มีได้หลาย Company")
add_bullet("Branch (สาขา) — Branch 00000 = สำนักงานใหญ่")
doc.add_paragraph()
add_para("ทุก Request ต้องมี AppContext (frozen dataclass):", size=10)
add_table(
    ["Field", "Type", "Description"],
    [
        ["firm_id", "int", "รหัสสำนักงานบัญชี"],
        ["company_id", "int", "รหัสบริษัทลูกค้า"],
        ["branch_id", "int", "รหัสสาขา"],
        ["user_id", "int", "รหัสผู้ใช้งาน"],
        ["user_role", "str", "firm_admin / accountant / junior / client_viewer / auditor"],
        ["period", "date | None", "งวดบัญชีที่ทำงาน"],
    ],
    [3, 4, 9]
)

add_heading("2.2 Database Isolation", 2)
add_table(
    ["Database File", "เนื้อหา"],
    [
        ["data/shared.sqlite", "Users, Firms, Companies, Branches (ข้อมูล platform กลาง)"],
        ["data/firm_{id}/company_{id}/db.sqlite", "ข้อมูลบัญชีทั้งหมดของบริษัทนั้น (COA, JE, AR, AP, FA ฯลฯ)"],
    ],
    [7, 9]
)

add_heading("2.3 Layer Diagram", 2)
layers = [
    ("Platform Layer", "Firms → Companies → Branches → Users → Roles → AppContext"),
    ("Core Layer", "COA → PostingEngine → JournalEntry → Ledger → Period → AccountBalance"),
    ("Module Layer", "AR · AP · Inventory · Fixed Assets · Bank · Tax · Payroll · Petty Cash · GL"),
    ("OCR Layer", "Claude Vision → Extractor → COA Classifier → Bank Reader"),
    ("Reports Layer", "Trial Balance · Balance Sheet · Income · Cashflow · Aging · Tax Report"),
    ("API Layer", "FastAPI Routers /api/v1/* — JSON responses"),
    ("UI Layer", "Jinja2 + HTMX + Alpine.js + Tailwind → HTML pages"),
]
add_table(["Layer", "Components"], layers, [4, 12])

# ════════════════════════════════════════════════════════════════════════════════
# 3. CORE ENGINE
# ════════════════════════════════════════════════════════════════════════════════
add_heading("3. Core Accounting Engine", 1)

add_heading("3.1 PostingEngine (Iron Rule)", 2)
add_para(
    "ทุก Journal Entry ต้องผ่าน PostingEngine เท่านั้น — ห้าม INSERT ตรงลงตาราง JournalEntry",
    bold=True, size=10, color=RGBColor(0xDC, 0x26, 0x26)
)
add_para("Signature: post(entry: JournalEntryInput, ctx: AppContext) → str (journal_no)", size=10)
add_para("Engine ตรวจสอบ Dr == Cr ก่อน post เสมอ แล้วเขียนพร้อมกันทั้ง:", size=10)
add_bullet("สมุดรายวัน (JournalEntry + JournalLine)")
add_bullet("บัญชีแยกประเภท (LedgerEntry)")
add_bullet("AccountBalance (pre-aggregated ต่อ period)")
doc.add_paragraph()

add_heading("3.2 Journal Types", 2)
add_table(
    ["Code", "ชื่อ", "ใช้กับ"],
    [
        ["GJ", "General Journal / สมุดรายวันทั่วไป", "Fixed Assets, Depreciation, Bank Transfer, Manual JE"],
        ["SJ", "Sales Journal / สมุดรายวันขาย", "AR — ออกใบแจ้งหนี้/ใบเสร็จ"],
        ["PJ", "Purchase Journal / สมุดรายวันซื้อ", "AP — บันทึกใบสั่งซื้อ/ใบแจ้งหนี้"],
        ["CR", "Cash Receipt / สมุดรับเงิน", "AR — รับชำระ"],
        ["CP", "Cash Payment / สมุดจ่ายเงิน", "AP — จ่ายชำระ, Petty Cash"],
    ],
    [1.5, 5.5, 9]
)

add_heading("3.3 COA Structure (NPAEs)", 2)
add_table(
    ["Category", "Code Range", "ตัวอย่างบัญชีสำคัญ"],
    [
        ["1 — สินทรัพย์", "1100–1299", "1101 เงินสด, 1102 ธนาคารกระแสรายวัน, 1110 ลูกหนี้การค้า, 1130 สินค้าคงเหลือ"],
        ["1 — สินทรัพย์ถาวร", "1220–1261", "1220 ที่ดิน, 1230/1231 อาคาร, 1240/1241 เครื่องจักร, 1260/1261 ยานพาหนะ"],
        ["2 — หนี้สิน", "2100–2199", "2101 เจ้าหนี้การค้า, 2102 เจ้าหนี้อื่น, 2103 เจ้าหนี้เช่าซื้อ, 2104 ดอกเบี้ยรอตัดบัญชี"],
        ["3 — ทุน", "3100–3299", "3101 ทุนชำระแล้ว, 3201 กำไรสะสม"],
        ["4 — รายได้", "4100–4299", "4101 ขายสินค้า, 4102 รายได้บริการ, 4201 ดอกเบี้ยรับ"],
        ["5 — ต้นทุน", "5100–5199", "5101 ต้นทุนสินค้าขาย, 5102 ต้นทุนบริการ"],
        ["6 — ค่าใช้จ่าย", "6100–6510", "6504 ค่าเสื่อมราคา, 6505 ค่าซ่อมแซม"],
        ["7 — ค่าใช้จ่ายทางการเงิน", "7100–7299", "7101 ดอกเบี้ยจ่าย, 7201 ภาษีนิติบุคคล"],
    ],
    [3.5, 3, 9.5]
)

# ════════════════════════════════════════════════════════════════════════════════
# 4. MODULES
# ════════════════════════════════════════════════════════════════════════════════
add_heading("4. โมดูลธุรกิจ (Business Modules)", 1)

add_heading("4.1 AR — Accounts Receivable (ลูกหนี้การค้า)", 2)
add_table(
    ["Feature", "รายละเอียด", "Journal"],
    [
        ["ใบเสนอราคา (Quotation)", "สร้าง/แปลงเป็น Invoice", "ไม่มี JE"],
        ["ใบแจ้งหนี้ขาย (Sales Invoice)", "ขายสด หรือ ขายเชื่อ, รองรับ VAT + WHT", "SJ: Dr ลูกหนี้/เงินสด | Cr รายได้"],
        ["ใบรับเงิน (Receipt)", "รับชำระบางส่วน/เต็มจำนวน", "CR: Dr ธนาคาร | Cr ลูกหนี้"],
        ["ใบวางบิล (Billing Note)", "รวม Invoice หลายใบส่งลูกค้า", "ไม่มี JE"],
        ["e-Tax Invoice", "ออกใบกำกับภาษีอิเล็กทรอนิกส์", ""],
        ["Aging Report", "วิเคราะห์อายุลูกหนี้", ""],
    ],
    [4, 9, 3]
)

add_heading("4.2 AP — Accounts Payable (เจ้าหนี้การค้า)", 2)
add_table(
    ["Feature", "รายละเอียด", "Journal"],
    [
        ["ใบสั่งซื้อ (Purchase Order)", "สร้าง PO ส่งผู้ขาย", "ไม่มี JE"],
        ["ใบรับสินค้า/บริการ (Purchase Invoice)", "ซื้อสด หรือ ซื้อเชื่อ, รองรับ VAT + WHT", "PJ: Dr สินค้า/ค่าใช้จ่าย | Cr เจ้าหนี้/เงินสด"],
        ["จ่ายชำระ (Payment)", "จ่ายบางส่วน/เต็มจำนวน", "CP: Dr เจ้าหนี้ | Cr ธนาคาร"],
        ["Aging Report", "วิเคราะห์อายุเจ้าหนี้", ""],
    ],
    [4, 9, 3]
)

add_heading("4.3 Inventory — สินค้าคงคลัง", 2)
add_table(
    ["Feature", "รายละเอียด"],
    [
        ["Product Master", "ทะเบียนสินค้า พร้อม 3 วิธีตีราคา: Average / FIFO / Specific"],
        ["รับสินค้า (Receive)", "Dr 1130 สินค้าคงเหลือ | Cr ตามแหล่งเงิน"],
        ["ตัดสินค้า (Issue)", "Dr 5101 ต้นทุนสินค้าขาย | Cr 1130"],
        ["เชื่อมต่อ AR/AP", "ซื้อผ่าน AP → receive_stock(post_journal=False) ป้องกัน double-post"],
        ["Stock Movement Report", "ประวัติการเคลื่อนไหวสต็อก"],
    ],
    [5, 11]
)

add_heading("4.4 Fixed Assets — สินทรัพย์ถาวร", 2)
add_para("โมดูลใหม่ล่าสุด ครบ 3 งาน:", size=10, bold=True)

add_heading("4.4.1 ทะเบียนสินทรัพย์", 3)
add_table(
    ["Field สำคัญ", "รายละเอียด"],
    [
        ["asset_code", "รหัสสินทรัพย์ (auto-generate, unique ต่อบริษัท)"],
        ["category", "land / building / equipment / vehicle / furniture / it"],
        ["asset_account + acc_depr_account", "COA ที่ผูกกับประเภท (อัตโนมัติ)"],
        ["cost / salvage_value", "ราคาทุน / มูลค่าซาก"],
        ["useful_life_months", "อายุการใช้งาน (เดือน)"],
        ["accumulated_depr / book_value", "ค่าเสื่อมสะสม / มูลค่าตามบัญชี (denormalized)"],
        ["status", "active / fully_depreciated / disposed"],
    ],
    [6, 10]
)

add_heading("4.4.2 แหล่งเงินทุน 4 ประเภท (Funding Type)", 3)
add_table(
    ["Funding Type", "Journal Entry ตอนซื้อ"],
    [
        ["💵 CASH_BANK (เงินสด/ธนาคาร)", "Dr สินทรัพย์ | Cr บัญชีธนาคาร (BankAccount master)"],
        ["👤 OWNER_CONTRIBUTION (เจ้าของลงทุน)", "Dr สินทรัพย์ | Cr 3101 ทุนชำระแล้ว"],
        ["📋 OTHER_PAYABLE (ซื้อเชื่อ)", "Dr สินทรัพย์ | Cr 2102 เจ้าหนี้อื่น"],
        ["📅 HIRE_PURCHASE (เช่าซื้อ)", "Dr สินทรัพย์ + Dr 2104 ดอกเบี้ยรอตัดบัญชี | Cr 2103 เจ้าหนี้เช่าซื้อ + Cr ดาวน์"],
    ],
    [6, 10]
)
add_para("เช่าซื้อ: สร้าง HirePurchaseInstallment schedule อัตโนมัติ — บันทึกผ่อนงวดได้ทีละงวด", size=10, indent=True)

add_heading("4.4.3 ค่าเสื่อมราคา — บัญชี vs ภาษี", 3)
add_table(
    ["Field", "รายละเอียด"],
    [
        ["book_useful_life_years", "อายุทางบัญชี (ผู้ใช้กำหนด, default = ค่าสรรพากร)"],
        ["book_monthly_depreciation", "ค่าเสื่อม/เดือน ทางบัญชี — ใช้ POST JE จริง"],
        ["tax_useful_life_years", "อายุทางภาษี (ขั้นต่ำตาม พ.ร.ฎ. 145/2527)"],
        ["tax_depreciable_cost", "ต้นทุนที่ใช้คำนวณภาษี (รถยนต์นั่ง cap ที่ 1,000,000 บาท)"],
        ["tax_monthly_depreciation", "ค่าเสื่อม/เดือน ทางภาษี — เก็บไว้อ้างอิง ไม่ post JE"],
    ],
    [5.5, 10.5]
)
add_table(
    ["ประเภทสินทรัพย์", "อัตราภาษีสูงสุด", "อายุขั้นต่ำ", "เพดานต้นทุน"],
    [
        ["ที่ดิน", "—", "—", "ไม่มีค่าเสื่อม"],
        ["อาคาร", "5%/ปี", "20 ปี", "—"],
        ["เครื่องจักร/ยานพาหนะ/เครื่องตกแต่ง", "20%/ปี", "5 ปี", "—"],
        ["คอมพิวเตอร์/IT", "33.33%/ปี", "3 ปี", "—"],
        ["รถยนต์นั่งส่วนบุคคล (≤10 ที่นั่ง)", "20%/ปี", "5 ปี", "1,000,000 บาท (มาตรา 5)"],
    ],
    [6, 3, 2.5, 4.5]
)
add_para("JE ค่าเสื่อม: Dr 6504 ค่าเสื่อมราคา | Cr บัญชีค่าเสื่อมสะสมของสินทรัพย์นั้น", size=10, indent=True)

add_heading("4.5 Bank — ธนาคาร", 2)
add_table(
    ["Feature", "รายละเอียด"],
    [
        ["BankAccount Master", "ทะเบียนบัญชีธนาคาร (ชื่อธนาคาร, เลขบัญชี, COA code, ประเภท: กระแสรายวัน/ออมทรัพย์/เงินสด)"],
        ["Bank Transfer", "โอนเงินระหว่างบัญชี → GJ: Dr ปลายทาง | Cr ต้นทาง"],
        ["Bank Reconciliation", "กระทบยอดบัญชีธนาคาร"],
        ["ใช้ใน AR/AP", "Dropdown บัญชีที่จ่าย/รับดึงจาก BankAccount master"],
        ["ใช้ใน FA", "แหล่งเงินทุน CASH_BANK และ HP down payment"],
    ],
    [5, 11]
)

add_heading("4.6 Tax — ภาษี", 2)
add_table(
    ["Feature", "รายละเอียด"],
    [
        ["VAT (ภาษีมูลค่าเพิ่ม)", "ภาษีซื้อ (1140) / ภาษีขาย (2120) คำนวณอัตโนมัติใน AR/AP"],
        ["WHT (ภาษีหัก ณ ที่จ่าย)", "คำนวณ WHT บน Invoice, ออกหนังสือรับรอง"],
        ["Tax Depreciation Report", "ตาราง book vs tax ค่าเสื่อม สำหรับ ภ.ง.ด.50"],
    ],
    [5, 11]
)

add_heading("4.7 โมดูลอื่น", 2)
add_table(
    ["Module", "รายละเอียด"],
    [
        ["Payroll (เงินเดือน)", "คำนวณเงินเดือน, SSF, ภาษี, post GJ อัตโนมัติ"],
        ["Petty Cash (เงินสดย่อย)", "เบิกจ่ายเงินสดย่อย, เติมเงิน, post CP"],
        ["GL (General Ledger)", "บันทึก JE ทั่วไป, Reversing Entry, Recurring Entry"],
        ["OCR", "อ่านใบแจ้งหนี้/Bank Statement ด้วย Claude Vision API → แนะนำบัญชี"],
    ],
    [4, 12]
)

# ════════════════════════════════════════════════════════════════════════════════
# 5. API ENDPOINTS
# ════════════════════════════════════════════════════════════════════════════════
add_heading("5. API Endpoints (/api/v1/)", 1)

add_table(
    ["Router File", "Prefix", "Endpoints หลัก"],
    [
        ["auth_routes.py", "/auth", "POST /login, POST /refresh, POST /logout"],
        ["coa_routes.py", "/coa", "GET /coa, POST /coa, PUT /coa/{code}"],
        ["journal_routes.py", "/journals", "GET /journals, POST /journals, GET /journals/{no}"],
        ["ledger_routes.py", "/ledger", "GET /ledger (บัญชีแยกประเภท)"],
        ["ar_routes.py", "/ar", "CRUD invoices, receipts, billing notes, quotations"],
        ["ap_routes.py", "/ap", "CRUD purchases, payments, purchase orders"],
        ["inv_routes.py", "/inv", "CRUD products, stock movements"],
        ["bank_routes.py", "/bank", "GET/POST /bank/accounts, POST /bank/transfers"],
        ["fa_routes.py", "/fa", "GET/POST /fa/assets, /fa/assets/{id}/installments, POST /fa/depreciation/run"],
        ["tax_routes.py", "/tax", "VAT report, WHT report"],
        ["report_routes.py", "/reports", "trial-balance, balance-sheet, income, aging (AR/AP)"],
        ["payroll_routes.py", "/payroll", "CRUD payroll runs"],
        ["ocr_routes.py", "/ocr", "POST /ocr/upload, GET /ocr/review/{id}"],
        ["platform_routes.py", "/platform", "Firms, Companies, Branches, Users CRUD"],
    ],
    [4, 3, 9]
)

add_para("Base URL: http://localhost:8000/api/v1/ (dev)  |  Auth: JWT Bearer token", size=9, color=GRAY)

# ════════════════════════════════════════════════════════════════════════════════
# 6. UI PAGES
# ════════════════════════════════════════════════════════════════════════════════
add_heading("6. หน้า UI (Web Pages)", 1)

add_table(
    ["URL Path", "หน้า"],
    [
        ["/", "Dashboard"],
        ["/journals", "สมุดรายวัน"],
        ["/ledger", "บัญชีแยกประเภท"],
        ["/ar/invoices", "รายการใบแจ้งหนี้ขาย"],
        ["/ar/invoices/new", "สร้างใบแจ้งหนี้ใหม่"],
        ["/ar/receipts/new", "บันทึกรับชำระ"],
        ["/ar/billing-notes", "ใบวางบิล"],
        ["/ar/quotations", "ใบเสนอราคา"],
        ["/ap/purchases", "รายการใบสั่งซื้อ/ใบแจ้งหนี้ซื้อ"],
        ["/ap/payments/new", "บันทึกจ่ายชำระ"],
        ["/inventory/products", "ทะเบียนสินค้า"],
        ["/inventory/movements", "ประวัติการเคลื่อนไหวสต็อก"],
        ["/bank/accounts", "ทะเบียนบัญชีธนาคาร"],
        ["/bank/transfers", "โอนเงินระหว่างบัญชี"],
        ["/assets", "ทะเบียนสินทรัพย์ถาวร"],
        ["/assets/{id}", "รายละเอียดสินทรัพย์ + ตารางผ่อนเช่าซื้อ"],
        ["/assets/tax-depreciation-report", "รายงานผลต่างค่าเสื่อม บัญชี vs ภาษี (ภ.ง.ด.50)"],
        ["/reports/trial-balance", "งบทดลอง"],
        ["/reports/balance-sheet", "งบดุล"],
        ["/reports/income", "งบกำไรขาดทุน"],
        ["/reports/aging/ar", "Aging ลูกหนี้"],
        ["/reports/aging/ap", "Aging เจ้าหนี้"],
        ["/ocr", "อัปโหลด OCR"],
        ["/firm", "Firm Dashboard"],
    ],
    [6, 10]
)

# ════════════════════════════════════════════════════════════════════════════════
# 7. DATABASE TABLES
# ════════════════════════════════════════════════════════════════════════════════
add_heading("7. ตาราง Database (Company DB)", 1)

add_table(
    ["Table", "Module", "รายละเอียด"],
    [
        ["chart_of_accounts", "Core", "ผังบัญชี 4 หลัก มาตรฐาน NPAEs"],
        ["periods", "Core", "งวดบัญชี (month/year)"],
        ["journal_entries + journal_lines", "Core", "สมุดรายวัน — ห้ามแก้โดยตรง"],
        ["ledger_entries", "Core", "บัญชีแยกประเภท"],
        ["account_balances", "Core", "ยอดต่อ account ต่อ period (pre-aggregated)"],
        ["ar_invoices + ar_invoice_lines", "AR", "ใบแจ้งหนี้ขาย"],
        ["ar_receipts", "AR", "ใบรับเงิน"],
        ["ar_billing_notes", "AR", "ใบวางบิล"],
        ["ar_quotations", "AR", "ใบเสนอราคา"],
        ["ap_purchases + ap_purchase_lines", "AP", "ใบรับสินค้า/ใบแจ้งหนี้ซื้อ"],
        ["ap_payments", "AP", "ใบจ่ายชำระ"],
        ["inv_products + inv_product_lots", "INV", "ทะเบียนสินค้า + lot ราคา"],
        ["inv_stock_movements", "INV", "ประวัติการเคลื่อนไหวสต็อก"],
        ["bank_accounts", "Bank", "ทะเบียนบัญชีธนาคาร"],
        ["bank_transfers", "Bank", "รายการโอนระหว่างบัญชี"],
        ["fa_assets", "FA", "ทะเบียนสินทรัพย์ถาวร"],
        ["fa_depreciation_records", "FA", "ประวัติการ post ค่าเสื่อมรายเดือน"],
        ["hp_installments", "FA", "ตาราง schedule ผ่อนชำระเช่าซื้อ"],
        ["contacts", "Shared", "ลูกค้า/ผู้ขาย (AR/AP)"],
    ],
    [5, 2.5, 8.5]
)

# ════════════════════════════════════════════════════════════════════════════════
# 8. IRON RULES
# ════════════════════════════════════════════════════════════════════════════════
add_heading("8. กฎเหล็ก (Iron Rules)", 1)
add_para("ห้ามละเมิดกฎต่อไปนี้ไม่ว่ากรณีใด:", bold=True, size=10)
rules = [
    "ทุก JournalEntry ต้องผ่าน PostingEngine เท่านั้น — ห้าม INSERT ตรง",
    "Dr รวม == Cr รวม เสมอ — PostingEngine ตรวจสอบก่อน post ทุกครั้ง",
    "ห้ามลบหรือแก้ไข JournalEntry โดยตรง — ใช้ Reversing Entry แทน",
    "ทุก entry ต้องมี branch_id, user_id, timestamp",
    "ข้อมูลแต่ละบริษัทแยกใน DB คนละไฟล์ — ห้ามปะปน",
    "ค่าเสื่อมที่ POST JE จริงต้องใช้ book_monthly_depreciation เสมอ (ไม่ใช่ tax)",
    "Inventory: ใช้ receive_stock(post_journal=False) เมื่อซื้อผ่าน AP เพื่อป้องกัน double-post",
]
for i, r in enumerate(rules, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{i}. {r}")
    run.font.size = Pt(10)
    if i <= 3:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

# ════════════════════════════════════════════════════════════════════════════════
# 9. DEVELOPMENT GUIDE
# ════════════════════════════════════════════════════════════════════════════════
add_heading("9. คู่มือการพัฒนา", 1)

add_heading("9.1 Environment Setup", 2)
cmds = [
    ("Clone & install", "git clone https://github.com/2bird555687-arch/accounting\ncd accounting\npython -m venv .venv\n.venv\\Scripts\\pip install -r requirements.txt"),
    ("Run dev server", ".venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --host 0.0.0.0 --port 8000"),
    ("Initialize DB", ".venv\\Scripts\\python.exe scripts\\init_company_db.py"),
    ("Migrate company DB", ".venv\\Scripts\\python.exe scripts\\migrate_company_db.py"),
    ("Seed COA + contacts", ".venv\\Scripts\\python.exe scripts\\seed_coa_contacts.py"),
    ("Seed bank accounts", ".venv\\Scripts\\python.exe scripts\\seed_bank_accounts.py"),
    ("Seed fixed assets", ".venv\\Scripts\\python.exe scripts\\seed_fixed_assets.py"),
]
for title, cmd in cmds:
    add_para(title, bold=True, size=10)
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)

add_heading("9.2 Adding a New Module", 2)
steps = [
    "สร้าง app/modules/{name}/models.py — SQLAlchemy 2.0 (Mapped[], mapped_column())",
    "สร้าง app/modules/{name}/schemas.py — Pydantic models",
    "สร้าง app/modules/{name}/service.py — business logic, ใช้ PostingEngine สำหรับ JE",
    "สร้าง app/api/v1/{name}_routes.py — FastAPI router",
    "Register router ใน app/api/v1/__init__.py",
    "เพิ่ม UI route ใน app/ui/routes.py",
    "สร้าง templates/{name}/*.html — extends base.html, ใช้ Alpine.js",
    "เพิ่มเมนูใน templates/base.html sidebar (sb-section, sb-group-head, sb-item)",
    "เพิ่ม CREATE TABLE ใน scripts/migrate_company_db.py",
]
for i, s in enumerate(steps, 1):
    add_bullet(f"{i}. {s}")

add_heading("9.3 Design System", 2)
add_table(
    ["Class / Element", "ใช้งาน"],
    [
        [".panel", "Card container ขาวพร้อม shadow"],
        [".panel-head", "Header ของ panel"],
        [".btn, .btn-primary, .btn-outline", "ปุ่มต่าง ๆ"],
        [".badge", "แสดงสถานะ (badge-green, badge-red, badge-yellow)"],
        [".sb-section", "หัวกลุ่มเมนู sidebar"],
        [".sb-group-head", "กลุ่มเมนูแบบ accordion"],
        [".sb-item", "รายการเมนูย่อย"],
        ["Alpine.js x-data / x-init", "State management ต่อ component"],
        ["fetch() + Bearer token", "เรียก API (token จาก cookie ผ่าน JS)"],
    ],
    [5.5, 10.5]
)

# ════════════════════════════════════════════════════════════════════════════════
# 10. RECENT CHANGES
# ════════════════════════════════════════════════════════════════════════════════
add_heading("10. ประวัติการพัฒนาล่าสุด", 1)
add_table(
    ["Commit", "งาน", "รายละเอียด"],
    [
        ["19ff8a1", "Tax Depreciation", "เพิ่ม default อายุการใช้งานตามสรรพากร + book/tax แยกกัน + รายงาน ภ.ง.ด.50"],
        ["d27598f", "FA Funding Types", "รองรับ 4 แหล่งเงินทุน: เงินสด/เจ้าของลงทุน/เจ้าหนี้อื่น/เช่าซื้อ + HP installments"],
        ["5c460e3", "Fixed Assets Module", "สร้างโมดูลสินทรัพย์ถาวรครบ: ทะเบียน, ค่าเสื่อมรายเดือน, UI"],
        ["7172a2f", "Bank Name Datalist", "เพิ่ม datalist 14 ธนาคารไทยในฟอร์มเพิ่มบัญชีธนาคาร"],
        ["1e43a2e", "Bank Module", "ทะเบียนบัญชีธนาคาร, โอนเงิน, AR/AP ใช้ BankAccount master"],
        ["d52bddd", "Inventory Module", "สินค้าคงคลัง 3 วิธีตีราคา, เชื่อมต่อ AR/AP"],
        ["70094b6", "Fix Reports 404", "แก้ Trial Balance / Balance Sheet คืน empty report แทน 404 เมื่อไม่มี Period"],
    ],
    [2.5, 5, 8.5]
)

# ════════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════════
out = r"D:\accounting\AccCloud_ProjectDoc.docx"
doc.save(out)
print(f"Saved: {out}")
